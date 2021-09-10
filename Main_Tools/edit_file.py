from docx import Document
from googletrans import Translator
from constants import TODO_SRC, COMPLETED_SRC
from constants import TUPLE_OF_RU_DICT, TUPLE_OF_EN_DICT


class EditFile:
    def __init__(self, translate_file):
        self.en_file = translate_file
        self.content = "\n"
        self.translated_ru_content = None

    def get_content_en_file(self):
        en_docx = Document(TODO_SRC + self.en_file)

        for para in en_docx.paragraphs:
            self.content += para.text + "\n"

        self.content = edit_by_dict(self.content, TUPLE_OF_EN_DICT)

    def translate_en_ru(self):
        translator = Translator(service_urls=["translate.googleapis.com"])
        self.translated_ru_content = translator.translate(self.content, src="en", dest="ru")

        self.translated_ru_content = edit_by_dict(self.translated_ru_content.text, TUPLE_OF_RU_DICT)

    def put_translated_file(self):
        new_doc = Document()

        new_doc.add_paragraph(self.translated_ru_content)

        if 'EN' in self.en_file:
            new_doc.save(COMPLETED_SRC + self.en_file.replace("EN", "RU"))
        else:
            new_doc.save(COMPLETED_SRC + self.en_file.replace(".docx", " RU.docx"))


def edit_by_dict(content, tuple_of_dict):
    if not tuple_of_dict:
        return print("There are no dictionaries")

    for dictionary in tuple_of_dict:
        if not dictionary:
            continue

        all_keys = dictionary.keys()
        for key in all_keys:
            if key in content:
                content = content.replace(key, dictionary[key])

        return content
